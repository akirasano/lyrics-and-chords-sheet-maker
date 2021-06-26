import docx
from docx import Document  # nopep
from docx.shared import Pt  # nopep
from docx.shared import Inches  # nopep
from docx.oxml.xmlchemy import ZeroOrOne  # nopep
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION


def si(o):
    print(type(o))
    print(o)
    print(dir(o))
    print()


def make(out, ctlist, song, artist, title_font_pt=14, font_pt=12, ruby_pt=12, ruby_offset=0, write_chord=True):

    document = Document()

    head_paragraph = document.add_paragraph(f'{song} / {artist}')
    head_paragraph.runs[0].font.size = Pt(title_font_pt)

    main_sec = document.add_section(WD_SECTION.CONTINUOUS)
    # ２段組
    # https://stackoverflow.com/questions/30707120/how-to-programatically-implement-columns-in-page-layout-as-of-in-ms-word-using-p
    sectPr = main_sec._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    def add_text_with_ruby(paragraph, text, rubytext):
        run = paragraph.add_run()

        def ruby_pr():
            hps = OxmlElement('w:hps')
            hps.set(qn('w:val'), str(ruby_pt * 2))

            raise_base = 32
            hraise = OxmlElement('w:hpsRaise')
            hraise.set(qn('w:val'),
                       str(raise_base + ruby_offset))

            # hbase = OxmlElement('w:hpsBaseText')
            # hbase.set(qn('w:val'), str(font_size))

            r = OxmlElement('w:rubyPr')
            r.append(hps)
            r.append(hraise)
            # r.append(hbase)

            return r

        def ruby_rt(rubytext):
            t = OxmlElement('w:t')
            t.text = rubytext

            rpr = OxmlElement('w:rPr')

            r = OxmlElement('w:r')
            r.append(rpr)
            r.append(t)

            rt = OxmlElement('w:rt')
            rt.append(r)

            return rt

        def rubyBase(text):
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_pt * 2))
            rpr = OxmlElement('w:rPr')
            rpr.append(sz)

            t = OxmlElement('w:t')
            t.text = text

            r = OxmlElement('w:r')
            r.append(rpr)
            r.append(t)

            rubyBase = OxmlElement('w:rubyBase')
            rubyBase.append(r)

            return rubyBase

        ruby = OxmlElement('w:ruby')
        ruby.append(ruby_pr())
        ruby.append(ruby_rt(rubytext))
        ruby.append(rubyBase(text))

        run._r.append(ruby)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = docx.shared.Pt(18)
    for line in ctlist:
        if len(line) == 0:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = docx.shared.Pt(18)
            continue
        else:
            if len(paragraph.runs) > 0:
                paragraph.add_run('\n')

        if write_chord:
            for p in line:
                chord_str = p[0]
                lyric0 = p[1][:1]
                lyric1 = p[1][1:]

                if lyric0 == '　' and len(lyric1) == 0:
                    lyric1 = '　'

                if chord_str != '':
                    chord_str = chord_str.replace('\\/', '/')
                    add_text_with_ruby(paragraph, lyric0, chord_str)
                else:
                    tr = paragraph.add_run(lyric0)
                    tr.font.size = Pt(font_pt)

                if len(lyric1) > 0:
                    tr = paragraph.add_run(lyric1)
                    tr.font.size = Pt(font_pt)
        else:
            for p in line:
                tr = paragraph.add_run(p[1])
                tr.font.size = Pt(font_pt)

    # margin setup
    c = 36000 # mm to margin
    margin = docx.shared.Mm(12.7) # mm
    for s in document.sections:
        s.top_margin = margin
        s.bottom_margin = margin
        s.left_margin = margin
        s.right_margin = margin

    document.save(out)
